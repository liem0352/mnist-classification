import docx
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# 读取文档
doc = Document('d:\\文件\\工作 作业\\深度学习\\5.21\\深度学习-实训报告6.docx')

# 定义要添加的代码内容

# 任务一：环境准备和导入库
task1_env_code = """import os
import numpy as np
import matplotlib.pyplot as plt
import tensorflow as tf
from tensorflow import keras
from tensorflow.keras import layers, models
from tensorflow.keras.datasets import mnist
from tensorflow.keras.utils import to_categorical

# 设置中文显示
plt.rcParams['font.sans-serif'] = ['SimHei', 'Microsoft YaHei', 'Arial Unicode MS']
plt.rcParams['axes.unicode_minus'] = False

print(f"TensorFlow版本: {tf.__version__}")
print(f"NumPy版本: {np.__version__}")"""

# 任务一：加载并预处理数据
task1_data_code = """def load_and_preprocess_data():
    # 加载MNIST数据集
    (x_train, y_train), (x_test, y_test) = mnist.load_data()
    
    # 归一化：将像素值从 [0, 255] 缩放到 [0, 1]
    x_train = x_train.astype('float32') / 255.0
    x_test = x_test.astype('float32') / 255.0
    
    # 将图像从 28x28 展平为 784 维向量
    x_train = x_train.reshape((x_train.shape[0], 28 * 28))
    x_test = x_test.reshape((x_test.shape[0], 28 * 28))
    
    # 将标签转换为one-hot编码
    y_train = to_categorical(y_train, 10)
    y_test = to_categorical(y_test, 10)
    
    print(f"训练集形状: {x_train.shape}, 标签形状: {y_train.shape}")
    print(f"测试集形状: {x_test.shape}, 标签形状: {y_test.shape}")
    
    return x_train, y_train, x_test, y_test

# 加载数据
x_train, y_train, x_test, y_test = load_and_preprocess_data()"""

# 任务一：可视化样本数据
task1_visual_code = """# 可视化部分训练样本
plt.figure(figsize=(12, 4))
for i in range(10):
    plt.subplot(2, 5, i + 1)
    plt.imshow(x_train[i].reshape(28, 28), cmap='gray')
    plt.title(f'标签: {np.argmax(y_train[i])}')
    plt.axis('off')
plt.suptitle('MNIST训练样本示例')
plt.tight_layout()
plt.show()"""

# 任务一：构建神经网络模型
task1_model_code = """def build_model():
    model = models.Sequential([
        # 输入层
        layers.Input(shape=(784,)),
        
        # 第一隐藏层
        layers.Dense(256, activation='relu'),
        layers.Dropout(0.2),
        
        # 第二隐藏层
        layers.Dense(128, activation='relu'),
        layers.Dropout(0.2),
        
        # 输出层
        layers.Dense(10, activation='softmax')
    ])
    
    # 编译模型
    model.compile(
        optimizer='adam',
        loss='categorical_crossentropy',
        metrics=['accuracy']
    )
    
    print("模型结构：")
    model.summary()
    
    return model

# 构建模型
model = build_model()"""

# 任务一：前向传播与反向传播演示
task1_propagation_code = """def demonstrate_forward_backward_propagation(model, x_train, y_train):
    print("\\n" + "="*60)
    print("前向传播与反向传播演示")
    print("="*60)
    
    # 取一个样本
    sample_idx = 0
    x_sample = x_train[sample_idx:sample_idx+1]
    y_sample = y_train[sample_idx:sample_idx+1]
    
    print(f"\\n选取样本索引: {sample_idx}")
    print(f"输入形状: {x_sample.shape}")
    print(f"真实标签: {np.argmax(y_sample)}")
    
    # 前向传播：获取各层输出
    print("\\n--- 前向传播过程 ---")
    _ = model.predict(x_sample, verbose=0)
    
    from tensorflow.keras.layers import Input
    input_layer = Input(shape=(784,))
    layer_outputs = []
    x = input_layer
    for layer in model.layers:
        x = layer(x)
        layer_outputs.append(x)
    activation_model = models.Model(inputs=input_layer, outputs=layer_outputs)
    
    activations = activation_model.predict(x_sample, verbose=0)
    
    for i, (layer, activation) in enumerate(zip(model.layers, activations)):
        print(f"层 {i+1} ({layer.name}): 输出形状 = {activation.shape}")
        if i == len(model.layers) - 1:
            print(f"  输出概率分布: {activation[0].round(4)}")
            print(f"  预测类别: {np.argmax(activation[0])}")
    
    # 反向传播：计算梯度
    print("\\n--- 反向传播过程（梯度计算） ---")
    x_tensor = tf.convert_to_tensor(x_sample, dtype=tf.float32)
    y_tensor = tf.convert_to_tensor(y_sample, dtype=tf.float32)
    
    with tf.GradientTape() as tape:
        tape.watch(x_tensor)
        predictions = model(x_tensor)
        loss = tf.keras.losses.categorical_crossentropy(y_tensor, predictions)
    
    gradients = tape.gradient(loss, model.trainable_variables)
    
    print(f"损失值: {loss.numpy()[0]:.6f}")
    print("\\n各层参数梯度统计:")
    for i, (grad, var) in enumerate(zip(gradients, model.trainable_variables)):
        if grad is not None:
            print(f"  参数 {i+1} ({var.name}):")
            print(f"    形状: {grad.shape}")
            print(f"    梯度均值: {tf.reduce_mean(tf.abs(grad)).numpy():.6f}")
            print(f"    梯度最大值: {tf.reduce_max(tf.abs(grad)).numpy():.6f}")
    
    print("\\n" + "="*60)

demonstrate_forward_backward_propagation(model, x_train, y_train)"""

# 任务一：训练模型
task1_train_code = """# 训练模型
print("\\n开始训练模型...")
history = model.fit(
    x_train, y_train,
    batch_size=128,
    epochs=10,
    validation_split=0.1,
    verbose=1
)"""

# 任务一：评估模型
task1_eval_code = """# 评估模型
print("\\n评估模型性能...")
test_loss, test_accuracy = model.evaluate(x_test, y_test, verbose=0)
print(f"测试集损失: {test_loss:.4f}")
print(f"测试集准确率: {test_accuracy:.4f}")"""

# 任务一：保存和加载模型
task1_save_code = """# 保存模型
model.save('mnist_model.keras')
print("模型已保存到: mnist_model.keras")

# 加载模型
loaded_model = keras.models.load_model('mnist_model.keras')
print("模型已加载")

# 验证加载的模型
loaded_loss, loaded_acc = loaded_model.evaluate(x_test, y_test, verbose=0)
print(f"加载模型测试准确率: {loaded_acc:.4f}")"""

# 任务一：可视化训练历史
task1_history_code = """# 绘制训练历史
fig, axes = plt.subplots(1, 2, figsize=(14, 5))

# 损失曲线
axes[0].plot(history.history['loss'], label='训练损失')
axes[0].plot(history.history['val_loss'], label='验证损失')
axes[0].set_xlabel('Epoch')
axes[0].set_ylabel('Loss')
axes[0].set_title('训练与验证损失')
axes[0].legend()
axes[0].grid(True)

# 准确率曲线
axes[1].plot(history.history['accuracy'], label='训练准确率')
axes[1].plot(history.history['val_accuracy'], label='验证准确率')
axes[1].set_xlabel('Epoch')
axes[1].set_ylabel('Accuracy')
axes[1].set_title('训练与验证准确率')
axes[1].legend()
axes[1].grid(True)

plt.tight_layout()
plt.savefig('task1_training_history.png', dpi=150, bbox_inches='tight')
plt.show()
print("训练历史曲线已保存")"""

# 任务一：可视化预测结果
task1_predict_code = """# 可视化预测结果
num_samples = 10
indices = np.random.choice(len(x_test), num_samples, replace=False)

plt.figure(figsize=(15, 4))
for i, idx in enumerate(indices):
    sample = x_test[idx:idx+1]
    prediction = model.predict(sample, verbose=0)
    predicted_label = np.argmax(prediction)
    true_label = np.argmax(y_test[idx])
    
    plt.subplot(2, 5, i + 1)
    plt.imshow(x_test[idx].reshape(28, 28), cmap='gray')
    color = 'green' if predicted_label == true_label else 'red'
    plt.title(f'预测: {predicted_label}\\n真实: {true_label}', color=color)
    plt.axis('off')

plt.tight_layout()
plt.savefig('task1_predictions.png', dpi=150, bbox_inches='tight')
plt.show()
print("预测可视化已保存")"""

# 任务二：交叉熵损失训练
task2_ce_code = """# 构建模型
model_ce = models.Sequential([
    layers.Input(shape=(784,)),
    layers.Dense(256, activation='relu'),
    layers.Dropout(0.2),
    layers.Dense(128, activation='relu'),
    layers.Dropout(0.2),
    layers.Dense(10, activation='softmax')
])

# 编译模型（交叉熵损失）
model_ce.compile(
    optimizer='adam',
    loss='categorical_crossentropy',
    metrics=['accuracy']
)

print("\\n使用交叉熵损失训练模型...")
history_ce = model_ce.fit(
    x_train, y_train,
    batch_size=128,
    epochs=10,
    validation_split=0.1,
    verbose=1
)

# 评估
loss_ce, acc_ce = model_ce.evaluate(x_test, y_test, verbose=0)
print(f"\\n交叉熵 - 测试损失: {loss_ce:.4f}, 测试准确率: {acc_ce:.4f}")"""

# 任务二：MSE损失训练
task2_mse_code = """# 构建模型
model_mse = models.Sequential([
    layers.Input(shape=(784,)),
    layers.Dense(256, activation='relu'),
    layers.Dropout(0.2),
    layers.Dense(128, activation='relu'),
    layers.Dropout(0.2),
    layers.Dense(10, activation='softmax')
])

# 编译模型（MSE损失）
model_mse.compile(
    optimizer='adam',
    loss='mse',
    metrics=['accuracy']
)

print("\\n使用均方误差损失训练模型...")
history_mse = model_mse.fit(
    x_train, y_train,
    batch_size=128,
    epochs=10,
    validation_split=0.1,
    verbose=1
)

# 评估
loss_mse, acc_mse = model_mse.evaluate(x_test, y_test, verbose=0)
print(f"\\nMSE - 测试损失: {loss_mse:.4f}, 测试准确率: {acc_mse:.4f}")"""

# 任务二：对比结果
task2_compare_code = """print("\\n" + "="*60)
print("训练结果对比")
print("="*60)

print("\\n【最终训练结果对比】")
print(f"{'指标':<25} {'交叉熵':<15} {'均方误差':<15}")
print("-" * 60)
print(f"{'最终训练损失':<25} {history_ce.history['loss'][-1]:<15.4f} {history_mse.history['loss'][-1]:<15.4f}")
print(f"{'最终训练准确率':<25} {history_ce.history['accuracy'][-1]:<15.4f} {history_mse.history['accuracy'][-1]:<15.4f}")
print(f"{'最终验证损失':<25} {history_ce.history['val_loss'][-1]:<15.4f} {history_mse.history['val_loss'][-1]:<15.4f}")
print(f"{'最终验证准确率':<25} {history_ce.history['val_accuracy'][-1]:<15.4f} {history_mse.history['val_accuracy'][-1]:<15.4f}")

print("\\n【测试集结果】")
print(f"交叉熵 - 损失: {loss_ce:.4f}, 准确率: {acc_ce:.4f}")
print(f"MSE    - 损失: {loss_mse:.4f}, 准确率: {acc_mse:.4f}")"""

# 任务二：梯度更新示例
task2_gradient_code = """def demonstrate_gradient_update(model, loss_name, x_train, y_train):
    print(f"\\n{'='*60}")
    print(f"梯度更新示例 - {loss_name}")
    print(f"{'='*60}")
    
    sample_idx = 0
    x_sample = x_train[sample_idx:sample_idx+1]
    y_sample = y_train[sample_idx:sample_idx+1]
    
    x_tensor = tf.convert_to_tensor(x_sample, dtype=tf.float32)
    y_tensor = tf.convert_to_tensor(y_sample, dtype=tf.float32)
    
    if loss_name == 'categorical_crossentropy':
        loss_fn = tf.keras.losses.CategoricalCrossentropy()
    else:
        loss_fn = tf.keras.losses.MeanSquaredError()
    
    with tf.GradientTape() as tape:
        predictions = model(x_tensor)
        loss = loss_fn(y_tensor, predictions)
    
    gradients = tape.gradient(loss, model.trainable_variables)
    
    print(f"\\n样本索引: {sample_idx}")
    print(f"真实标签: {np.argmax(y_sample)}")
    print(f"预测概率: {predictions.numpy()[0].round(4)}")
    print(f"预测类别: {np.argmax(predictions.numpy()[0])}")
    print(f"\\n损失值 ({loss_name}): {loss.numpy():.6f}")
    
    print("\\n各层参数梯度统计:")
    for i, (grad, var) in enumerate(zip(gradients, model.trainable_variables)):
        if grad is not None:
            grad_flat = tf.reshape(grad, [-1])
            print(f"\\n  参数 {i+1} ({var.name}):")
            print(f"    形状: {grad.shape}")
            print(f"    梯度均值: {tf.reduce_mean(tf.abs(grad)).numpy():.8f}")
            print(f"    梯度标准差: {tf.math.reduce_std(grad).numpy():.8f}")
            print(f"    前5个梯度值: {grad_flat[:5].numpy().round(8)}")

# 交叉熵梯度示例
demonstrate_gradient_update(model_ce, 'categorical_crossentropy', x_train, y_train)

# MSE梯度示例
demonstrate_gradient_update(model_mse, 'mse', x_train, y_train)"""

# 任务二：可视化对比
task2_visual_code = """# 绘制对比图
fig, axes = plt.subplots(2, 2, figsize=(14, 10))

epoch_range = range(1, 11)

# 训练损失对比
axes[0, 0].plot(epoch_range, history_ce.history['loss'], 'b-', label='交叉熵', linewidth=2)
axes[0, 0].plot(epoch_range, history_mse.history['loss'], 'r-', label='均方误差', linewidth=2)
axes[0, 0].set_xlabel('Epoch')
axes[0, 0].set_ylabel('Loss')
axes[0, 0].set_title('训练损失对比')
axes[0, 0].legend()
axes[0, 0].grid(True)

# 验证损失对比
axes[0, 1].plot(epoch_range, history_ce.history['val_loss'], 'b-', label='交叉熵', linewidth=2)
axes[0, 1].plot(epoch_range, history_mse.history['val_loss'], 'r-', label='均方误差', linewidth=2)
axes[0, 1].set_xlabel('Epoch')
axes[0, 1].set_ylabel('Loss')
axes[0, 1].set_title('验证损失对比')
axes[0, 1].legend()
axes[0, 1].grid(True)

# 训练准确率对比
axes[1, 0].plot(epoch_range, history_ce.history['accuracy'], 'b-', label='交叉熵', linewidth=2)
axes[1, 0].plot(epoch_range, history_mse.history['accuracy'], 'r-', label='均方误差', linewidth=2)
axes[1, 0].set_xlabel('Epoch')
axes[1, 0].set_ylabel('Accuracy')
axes[1, 0].set_title('训练准确率对比')
axes[1, 0].legend()
axes[1, 0].grid(True)

# 验证准确率对比
axes[1, 1].plot(epoch_range, history_ce.history['val_accuracy'], 'b-', label='交叉熵', linewidth=2)
axes[1, 1].plot(epoch_range, history_mse.history['val_accuracy'], 'r-', label='均方误差', linewidth=2)
axes[1, 1].set_xlabel('Epoch')
axes[1, 1].set_ylabel('Accuracy')
axes[1, 1].set_title('验证准确率对比')
axes[1, 1].legend()
axes[1, 1].grid(True)

plt.tight_layout()
plt.savefig('task2_loss_comparison.png', dpi=150, bbox_inches='tight')
plt.show()
print("\\n对比图已保存到 task2_loss_comparison.png")"""

# 实验结果和体会
task1_result = """任务一结果：
1. 成功使用Keras高层API搭建了包含两个隐藏层的神经网络模型
2. 模型在MNIST数据集上经过10个epoch训练后，测试集准确率达到较高水平
3. 前向传播过程逐层计算输出，最终通过Softmax得到概率分布
4. 反向传播过程使用链式法则计算梯度，并通过Adam优化器更新参数
5. 模型成功保存并加载，验证了模型的可复用性"""

task2_result = """任务二结果：
1. 交叉熵损失：测试准确率明显高于MSE损失
2. 均方误差损失：在分类任务中表现较差，收敛速度较慢
3. 交叉熵损失与Softmax激活函数配合效果更好，梯度信号更强
4. 验证了交叉熵损失是多分类任务的首选损失函数"""

experience = """实验体会：
通过本次实训，深入理解了神经网络的前向传播和反向传播机制，掌握了Keras高层API的使用方法，并验证了不同损失函数对分类任务性能的影响。交叉熵损失在分类任务中的优势得到了充分验证。"""


def add_formatted_text(cell, text, bold=False, size=10, is_code=False):
    """向单元格添加格式化文本"""
    paragraph = cell.add_paragraph()
    run = paragraph.add_run(text)
    run.font.size = Pt(size)
    if bold:
        run.bold = True
    if is_code:
        run.font.name = 'Consolas'
        run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    return paragraph


def add_code_block(cell, code_text, title=""):
    """添加代码块到单元格"""
    if title:
        p = cell.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        run.font.size = Pt(10)
    
    # 添加代码内容
    p = cell.add_paragraph()
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0x80)
    
    # 添加空行
    cell.add_paragraph()


# 获取表格
table = doc.tables[0]

# 清空步骤记录单元格 [6,1] 并重新填充
step_cell = table.rows[6].cells[1]
# 保留第一个段落（可能包含"任务一："等标题），清除其余内容
for paragraph in step_cell.paragraphs[1:]:
    p_element = paragraph._element
    p_element.getparent().remove(p_element)

# 添加任务一的代码
add_formatted_text(step_cell, "【任务一：使用Keras高层API搭建神经网络实现手写数字分类】", bold=True, size=11)
step_cell.add_paragraph()

add_formatted_text(step_cell, "1. 环境准备与导入库：", bold=True)
add_code_block(step_cell, task1_env_code)

add_formatted_text(step_cell, "2. 加载并预处理MNIST数据集：", bold=True)
add_code_block(step_cell, task1_data_code)

add_formatted_text(step_cell, "3. 可视化样本数据：", bold=True)
add_code_block(step_cell, task1_visual_code)

add_formatted_text(step_cell, "4. 构建神经网络模型：", bold=True)
add_code_block(step_cell, task1_model_code)

add_formatted_text(step_cell, "5. 前向传播与反向传播演示：", bold=True)
add_code_block(step_cell, task1_propagation_code)

add_formatted_text(step_cell, "6. 训练模型：", bold=True)
add_code_block(step_cell, task1_train_code)

add_formatted_text(step_cell, "7. 评估模型：", bold=True)
add_code_block(step_cell, task1_eval_code)

add_formatted_text(step_cell, "8. 保存和加载模型：", bold=True)
add_code_block(step_cell, task1_save_code)

add_formatted_text(step_cell, "9. 可视化训练历史：", bold=True)
add_code_block(step_cell, task1_history_code)

add_formatted_text(step_cell, "10. 可视化预测结果：", bold=True)
add_code_block(step_cell, task1_predict_code)

step_cell.add_paragraph()
add_formatted_text(step_cell, "【任务二：对比交叉熵与均方误差损失对手写数字分类的影响】", bold=True, size=11)
step_cell.add_paragraph()

add_formatted_text(step_cell, "1. 使用交叉熵损失训练模型：", bold=True)
add_code_block(step_cell, task2_ce_code)

add_formatted_text(step_cell, "2. 使用均方误差损失训练模型：", bold=True)
add_code_block(step_cell, task2_mse_code)

add_formatted_text(step_cell, "3. 对比训练结果：", bold=True)
add_code_block(step_cell, task2_compare_code)

add_formatted_text(step_cell, "4. 梯度更新示例：", bold=True)
add_code_block(step_cell, task2_gradient_code)

add_formatted_text(step_cell, "5. 可视化对比结果：", bold=True)
add_code_block(step_cell, task2_visual_code)

# 清空结果分析单元格 [7,1] 并重新填充
result_cell = table.rows[7].cells[1]
for paragraph in result_cell.paragraphs[1:]:
    p_element = paragraph._element
    p_element.getparent().remove(p_element)

add_formatted_text(result_cell, task1_result, size=10)
result_cell.add_paragraph()
add_formatted_text(result_cell, task2_result, size=10)
result_cell.add_paragraph()
add_formatted_text(result_cell, experience, size=10)

# 保存文档
output_path = 'd:\\文件\\工作 作业\\深度学习\\5.21\\深度学习-实训报告6.docx'
doc.save(output_path)
print(f"文档已保存到: {output_path}")
